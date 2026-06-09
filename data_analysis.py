#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
懂球网比赛数据分析脚本
"""

import pandas as pd
from datetime import datetime

def load_data():
    basic = pd.read_csv('basic.csv', encoding='utf-8-sig')
    head_to_head = pd.read_csv('head_to_head.csv', encoding='utf-8-sig')
    home_recent = pd.read_csv('home_recent.csv', encoding='utf-8-sig')
    away_recent = pd.read_csv('away_recent.csv', encoding='utf-8-sig')
    
    head_to_head['日期'] = pd.to_datetime(head_to_head['日期'])
    home_recent['日期'] = pd.to_datetime(home_recent['日期'])
    away_recent['日期'] = pd.to_datetime(away_recent['日期'])
    
    head_to_head[['主队得分', '客队得分']] = head_to_head['比分'].str.split('-', expand=True).astype(int)
    home_recent[['主队得分', '客队得分']] = home_recent['比分'].str.split('-', expand=True).astype(int)
    away_recent[['主队得分', '客队得分']] = away_recent['比分'].str.split('-', expand=True).astype(int)
    
    return basic, head_to_head, home_recent, away_recent

def analyze_head_to_head(head_to_head, home_team, away_team):
    analysis = {}
    analysis['总交锋次数'] = len(head_to_head)
    
    home_wins = head_to_head[(head_to_head['主队'] == home_team) & (head_to_head['主队得分'] > head_to_head['客队得分'])].shape[0]
    home_wins += head_to_head[(head_to_head['客队'] == home_team) & (head_to_head['客队得分'] > head_to_head['主队得分'])].shape[0]
    
    away_wins = head_to_head[(head_to_head['主队'] == away_team) & (head_to_head['主队得分'] > head_to_head['客队得分'])].shape[0]
    away_wins += head_to_head[(head_to_head['客队'] == away_team) & (head_to_head['客队得分'] > head_to_head['主队得分'])].shape[0]
    
    analysis['主队胜场'] = home_wins
    analysis['客队胜场'] = away_wins
    analysis['主队胜率'] = f"{(home_wins / len(head_to_head)) * 100:.1f}%"
    analysis['客队胜率'] = f"{(away_wins / len(head_to_head)) * 100:.1f}%"
    
    analysis['主队场均得分'] = f"{head_to_head.apply(lambda x: x['主队得分'] if x['主队'] == home_team else x['客队得分'], axis=1).mean():.1f}"
    analysis['客队场均得分'] = f"{head_to_head.apply(lambda x: x['主队得分'] if x['主队'] == away_team else x['客队得分'], axis=1).mean():.1f}"
    
    recent_5 = head_to_head.sort_values('日期', ascending=False).head(5)
    analysis['最近5次交锋'] = []
    for _, row in recent_5.iterrows():
        analysis['最近5次交锋'].append({
            '日期': row['日期'].strftime('%Y-%m-%d'),
            '主队': row['主队'],
            '客队': row['客队'],
            '比分': row['比分'],
            '胜者': row['主队'] if row['主队得分'] > row['客队得分'] else row['客队']
        })
    
    return analysis

def analyze_recent_form(recent_data, team_name):
    analysis = {}
    analysis['比赛总数'] = len(recent_data)
    
    team_wins = recent_data[(recent_data['主队'] == team_name) & (recent_data['主队得分'] > recent_data['客队得分'])].shape[0]
    team_wins += recent_data[(recent_data['客队'] == team_name) & (recent_data['客队得分'] > recent_data['主队得分'])].shape[0]
    analysis['胜场'] = team_wins
    analysis['负场'] = len(recent_data) - team_wins
    analysis['胜率'] = f"{(team_wins / len(recent_data)) * 100:.1f}%"
    
    home_games = recent_data[recent_data['主队'] == team_name]
    if len(home_games) > 0:
        home_wins = home_games[home_games['主队得分'] > home_games['客队得分']].shape[0]
        analysis['主场比赛数'] = len(home_games)
        analysis['主场胜率'] = f"{(home_wins / len(home_games)) * 100:.1f}%"
        analysis['主场场均得分'] = f"{home_games['主队得分'].mean():.1f}"
        analysis['主场场均失'] = f"{home_games['客队得分'].mean():.1f}"
    
    away_games = recent_data[recent_data['客队'] == team_name]
    if len(away_games) > 0:
        away_wins = away_games[away_games['客队得分'] > away_games['主队得分']].shape[0]
        analysis['客场比赛数'] = len(away_games)
        analysis['客场胜率'] = f"{(away_wins / len(away_games)) * 100:.1f}%"
        analysis['客场场均得分'] = f"{away_games['客队得分'].mean():.1f}"
        analysis['客场场均失'] = f"{away_games['主队得分'].mean():.1f}"
    
    team_scores = recent_data.apply(lambda x: x['主队得分'] if x['主队'] == team_name else x['客队得分'], axis=1)
    opp_scores = recent_data.apply(lambda x: x['客队得分'] if x['主队'] == team_name else x['主队得分'], axis=1)
    analysis['场均得分'] = f"{team_scores.mean():.1f}"
    analysis['场均失分'] = f"{opp_scores.mean():.1f}"
    analysis['净胜分'] = f"{(team_scores - opp_scores).mean():.1f}"
    
    results = []
    for _, row in recent_data.sort_values('日期', ascending=True).iterrows():
        if (row['主队'] == team_name and row['主队得分'] > row['客队得分']) or \
           (row['客队'] == team_name and row['客队得分'] > row['主队得分']):
            results.append('W')
        else:
            results.append('L')
    
    max_win_streak = 0
    current_streak = 0
    for r in results:
        if r == 'W':
            current_streak += 1
            max_win_streak = max(max_win_streak, current_streak)
        else:
            current_streak = 0
    analysis['最长连胜'] = max_win_streak
    
    max_lose_streak = 0
    current_streak = 0
    for r in results:
        if r == 'L':
            current_streak += 1
            max_lose_streak = max(max_lose_streak, current_streak)
        else:
            current_streak = 0
    analysis['最长连败'] = max_lose_streak
    
    recent_10 = recent_data.sort_values('日期', ascending=False).head(10)
    analysis['最近10场'] = []
    for _, row in recent_10.iterrows():
        analysis['最近10场'].append({
            '日期': row['日期'].strftime('%Y-%m-%d'),
            '赛事': row['赛事'],
            '对手': row['客队'] if row['主队'] == team_name else row['主队'],
            '结果': '胜' if ((row['主队'] == team_name and row['主队得分'] > row['客队得分']) or 
                          (row['客队'] == team_name and row['客队得分'] > row['主队得分'])) else '负',
            '比分': row['比分']
        })
    
    return analysis

def generate_report(basic, h2h_analysis, home_analysis, away_analysis):
    home_team = basic['主队'][0]
    away_team = basic['客队'][0]
    match_date = basic['日期'][0]
    
    report = f"""# 克利夫兰骑士 vs 底特律活塞 比赛数据分析报告

## 一、比赛基本信息

| 项目 | 内容 |
|------|------|
| 比赛ID | {basic['比赛ID'][0]} |
| 主队 | {home_team} |
| 客队 | {away_team} |
| 比赛日期 | {match_date} |

---

## 二、双方交锋历史分析

### 2.1 交锋概况

| 项目 | {home_team} | {away_team} |
|------|-------------|-------------|
| 总交锋次数 | {h2h_analysis['总交锋次数']} 次 | - |
| 胜场数 | {h2h_analysis['主队胜场']} 场 | {h2h_analysis['客队胜场']} 场 |
| 胜率 | {h2h_analysis['主队胜率']} | {h2h_analysis['客队胜率']} |
| 场均得分 | {h2h_analysis['主队场均得分']} 分 | {h2h_analysis['客队场均得分']} 分 |

### 2.2 交锋历史趋势

从历史交锋数据来看，{home_team} 在双方对决中占据明显优势，胜率达到 {h2h_analysis['主队胜率']}，
而 {away_team} 的胜率为 {h2h_analysis['客队胜率']}。

### 2.3 最近5次交锋记录

| 日期 | 主队 | 客队 | 比分 | 胜者 |
|------|------|------|------|------|"""
    
    for record in h2h_analysis['最近5次交锋']:
        report += f"""
| {record['日期']} | {record['主队']} | {record['客队']} | {record['比分']} | {record['胜者']} |"""
    
    report += f"""

---

## 三、{home_team} 近期战绩分析

### 3.1 整体表现

| 项目 | 数据 |
|------|------|
| 近期比赛数 | {home_analysis['比赛总数']} 场 |
| 胜场 | {home_analysis['胜场']} 场 |
| 负场 | {home_analysis['负场']} 场 |
| 胜率 | {home_analysis['胜率']} |
| 场均得分 | {home_analysis['场均得分']} 分 |
| 场均失分 | {home_analysis['场均失分']} 分 |
| 净胜分 | {home_analysis['净胜分']} 分 |
| 最长连胜 | {home_analysis['最长连胜']} 场 |
| 最长连败 | {home_analysis['最长连败']} 场 |

### 3.2 主客场表现

| 项目 | 主场 | 客场 |
|------|------|------|
| 比赛数 | {home_analysis['主场比赛数']} 场 | {home_analysis['客场比赛数']} 场 |
| 胜率 | {home_analysis['主场胜率']} | {home_analysis['客场胜率']} |
| 场均得分 | {home_analysis['主场场均得分']} 分 | {home_analysis['客场场均得分']} 分 |
| 场均失分 | {home_analysis['主场场均失']} 分 | {home_analysis['客场场均失']} 分 |

### 3.3 最近10场比赛记录

| 日期 | 对手 | 结果 | 比分 |
|------|------|------|------|"""
    
    for record in home_analysis['最近10场']:
        report += f"""
| {record['日期']} | {record['对手']} | {record['结果']} | {record['比分']} |"""
    
    report += f"""

---

## 四、{away_team} 近期战绩分析

### 4.1 整体表现

| 项目 | 数据 |
|------|------|
| 近期比赛数 | {away_analysis['比赛总数']} 场 |
| 胜场 | {away_analysis['胜场']} 场 |
| 负场 | {away_analysis['负场']} 场 |
| 胜率 | {away_analysis['胜率']} |
| 场均得分 | {away_analysis['场均得分']} 分 |
| 场均失分 | {away_analysis['场均失分']} 分 |
| 净胜分 | {away_analysis['净胜分']} 分 |
| 最长连胜 | {away_analysis['最长连胜']} 场 |
| 最长连败 | {away_analysis['最长连败']} 场 |

### 4.2 主客场表现

| 项目 | 主场 | 客场 |
|------|------|------|
| 比赛数 | {away_analysis['主场比赛数']} 场 | {away_analysis['客场比赛数']} 场 |
| 胜率 | {away_analysis['主场胜率']} | {away_analysis['客场胜率']} |
| 场均得分 | {away_analysis['主场场均得分']} 分 | {away_analysis['客场场均得分']} 分 |
| 场均失分 | {away_analysis['主场场均失']} 分 | {away_analysis['客场场均失']} 分 |

### 4.3 最近10场比赛记录

| 日期 | 对手 | 结果 | 比分 |
|------|------|------|------|"""
    
    for record in away_analysis['最近10场']:
        report += f"""
| {record['日期']} | {record['对手']} | {record['结果']} | {record['比分']} |"""
    
    report += f"""

---

## 五、双方对比分析

### 5.1 综合对比

| 项目 | {home_team} | {away_team} |
|------|-------------|-------------|
| 近期胜率 | {home_analysis['胜率']} | {away_analysis['胜率']} |
| 场均得分 | {home_analysis['场均得分']} 分 | {away_analysis['场均得分']} 分 |
| 场均失分 | {home_analysis['场均失分']} 分 | {away_analysis['场均失分']} 分 |
| 净胜分 | {home_analysis['净胜分']} 分 | {away_analysis['净胜分']} 分 |
| 交锋胜率 | {h2h_analysis['主队胜率']} | {h2h_analysis['客队胜率']} |

### 5.2 数据分析结论

1. 交锋历史优势: {home_team} 在历史交锋中占据明显优势，胜率达到 {h2h_analysis['主队胜率']}，
   近期5次交锋也保持着不错的战绩。

2. 近期状态对比: 
   - {home_team} 近期胜率为 {home_analysis['胜率']}，净胜分 {home_analysis['净胜分']}，表现较为稳定。
   - {away_team} 近期胜率为 {away_analysis['胜率']}，净胜分 {away_analysis['净胜分']}，状态同样出色。

3. 主客场表现:
   - {home_team} 主场胜率 {home_analysis['主场胜率']}，客场胜率 {home_analysis['客场胜率']}，
     主场表现更具优势。
   - {away_team} 主场胜率 {away_analysis['主场胜率']}，客场胜率 {away_analysis['客场胜率']}，
     客战能力较强。

### 5.3 比赛预测

综合以上分析，本场比赛 {home_team} 占据以下优势：
- 历史交锋心理优势
- 主场作战优势
- 近期状态稳定

预计 {home_team} 获胜概率较大，但 {away_team} 近期状态出色，也具备爆冷的可能性。

---

## 六、数据说明

- 数据来源: 懂球网 (dongqiudi.com)
- 采集时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}
- 比赛状态: 本场比赛（{match_date}）数据尚未更新完整

---

报告结束
"""
    
    return report

def convert_to_docx(md_file, docx_file):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_table = False
    table_data = []
    
    for line in lines:
        line = line.rstrip('\n')
        
        if line.startswith('# '):
            heading = line[2:]
            paragraph = doc.add_heading(heading, level=1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0]
            run.font.name = '微软雅黑'
            run.font.size = Pt(16)
            run.font.bold = True
            doc.add_paragraph()
            
        elif line.startswith('## '):
            heading = line[3:]
            paragraph = doc.add_heading(heading, level=2)
            run = paragraph.runs[0]
            run.font.name = '微软雅黑'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
            
        elif line.startswith('### '):
            heading = line[4:]
            paragraph = doc.add_heading(heading, level=3)
            run = paragraph.runs[0]
            run.font.name = '微软雅黑'
            run.font.size = Pt(12)
            run.font.bold = True
            
        elif line.startswith('|') and '---' in line:
            in_table = True
            table_data = []
            
        elif line.startswith('|') and in_table:
            cells = [c.strip() for c in line.split('|')[1:-1]]
            table_data.append(cells)
            
        elif in_table and not line.startswith('|'):
            if table_data:
                rows = len(table_data)
                cols = len(table_data[0])
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                for i, row_data in enumerate(table_data):
                    for j, cell_text in enumerate(row_data):
                        cell = table.cell(i, j)
                        paragraph = cell.paragraphs[0]
                        paragraph.text = cell_text
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        run = paragraph.runs[0]
                        run.font.name = '微软雅黑'
                        run.font.size = Pt(10)
                
                doc.add_paragraph()
            
            in_table = False
            table_data = []
            
        elif line.strip() == '':
            doc.add_paragraph()
            
        else:
            paragraph = doc.add_paragraph(line)
            paragraph.style = 'No Spacing'
    
    doc.save(docx_file)

def main():
    basic, head_to_head, home_recent, away_recent = load_data()
    home_team = basic['主队'][0]
    away_team = basic['客队'][0]
    
    h2h_analysis = analyze_head_to_head(head_to_head, home_team, away_team)
    home_analysis = analyze_recent_form(home_recent, home_team)
    away_analysis = analyze_recent_form(away_recent, away_team)
    
    report = generate_report(basic, h2h_analysis, home_analysis, away_analysis)
    
    with open('analysis_report.md', 'w', encoding='utf-8') as f:
        f.write(report)
    print("分析报告 (Markdown) 已保存: analysis_report.md")
    
    convert_to_docx('analysis_report.md', 'analysis_report.docx')
    print("分析报告 (Word) 已保存: analysis_report.docx")

if __name__ == '__main__':
    main()
