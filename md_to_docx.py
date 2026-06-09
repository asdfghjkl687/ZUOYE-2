#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将Markdown格式的实验报告转换为Word文档
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import re

def md_to_docx(md_file, docx_file):
    """
    将Markdown文件转换为Word文档
    
    Args:
        md_file (str): 输入的Markdown文件路径
        docx_file (str): 输出的Word文件路径
    """
    # 创建文档对象
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # 用于跟踪当前是否在表格中
    in_table = False
    table_data = []
    
    for line in lines:
        line = line.rstrip('\n')
        
        # 处理标题
        if line.startswith('# '):
            # 一级标题
            heading = line[2:]
            paragraph = doc.add_heading(heading, level=1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0]
            run.font.name = '微软雅黑'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            doc.add_paragraph()  # 添加空行
            
        elif line.startswith('## '):
            # 二级标题
            heading = line[3:]
            paragraph = doc.add_heading(heading, level=2)
            run = paragraph.runs[0]
            run.font.name = '微软雅黑'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
            
        elif line.startswith('### '):
            # 三级标题
            heading = line[4:]
            paragraph = doc.add_heading(heading, level=3)
            run = paragraph.runs[0]
            run.font.name = '微软雅黑'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
        elif line.startswith('#### '):
            # 四级标题
            heading = line[5:]
            paragraph = doc.add_heading(heading, level=4)
            run = paragraph.runs[0]
            run.font.name = '微软雅黑'
            run.font.size = Pt(11)
            run.font.bold = True
            
        elif line.startswith('|') and '---' in line:
            # 表格分隔线，开始处理表格
            in_table = True
            table_data = []
            
        elif line.startswith('|') and in_table:
            # 表格内容行
            # 去除首尾的|，按|分割
            cells = [c.strip() for c in line.split('|')[1:-1]]
            table_data.append(cells)
            
        elif in_table and not line.startswith('|'):
            # 表格结束
            if table_data:
                # 创建表格
                rows = len(table_data)
                cols = len(table_data[0])
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                # 填充表格内容
                for i, row_data in enumerate(table_data):
                    for j, cell_text in enumerate(row_data):
                        cell = table.cell(i, j)
                        paragraph = cell.paragraphs[0]
                        paragraph.text = cell_text
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        
                        # 设置字体
                        run = paragraph.runs[0]
                        run.font.name = '微软雅黑'
                        run.font.size = Pt(10)
                
                # 添加空行
                doc.add_paragraph()
            
            in_table = False
            table_data = []
            
        elif line.startswith('```'):
            # 代码块开始/结束
            # 简单处理：添加段落标记
            paragraph = doc.add_paragraph('【代码块】')
            paragraph.style = 'No Spacing'
            
        elif line.startswith('- ') or line.startswith('1. ') or line.startswith('* '):
            # 列表项
            text = line[2:]
            paragraph = doc.add_paragraph(text, style='List Bullet')
            
        elif line.startswith('**') and line.endswith('**'):
            # 粗体文本
            text = line[2:-2]
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.bold = True
            
        elif line.strip() == '':
            # 空行
            doc.add_paragraph()
            
        else:
            # 普通文本段落
            paragraph = doc.add_paragraph(line)
            paragraph.style = 'No Spacing'
    
    # 保存文档
    doc.save(docx_file)
    print(f"Word文档 {docx_file} 已保存成功！")

if __name__ == '__main__':
    import sys
    
    if len(sys.argv) == 3:
        md_to_docx(sys.argv[1], sys.argv[2])
    else:
        md_to_docx('experiment_report.md', 'experiment_report.docx')
